from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
REPORT_DIR = ROOT / "reports" / "esg_chunk_report_esg_3b1c0196c0f9"
DATA_PATH = REPORT_DIR / "query_results.json"
VALIDATION_PATH = (
    ROOT
    / "outputs"
    / "esg_chunk_handoff_2000_esg_3b1c0196c0f9"
    / "VALIDATION.json"
)
STATS_PATH = REPORT_DIR / "chunk_stats.json"
OUT_PATH = REPORT_DIR / "ESG_Chunk_Dataset_and_2000_Chunk_Handoff.docx"
CHART_PATH = REPORT_DIR / "topic_distribution.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
DARK = "202124"
GREEN = "2E7D32"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "w:" + edge
        node = tc_mar.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_text(paragraph, text, *, size=10.5, bold=False, color=DARK, align=None) -> None:
    paragraph.clear()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_run_font(run, size=9, color=MID_GRAY)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_header_footer(section) -> None:
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.4)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_text(header, "ESG CHUNK DATASET • TECHNICAL REPORT", size=8.5, bold=True, color=MID_GRAY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_text(footer, "Dataset esg_3b1c0196c0f9  •  Page ", size=9, color=MID_GRAY)
    add_page_field(footer)


def make_chart(topic_rows: list[dict]) -> None:
    rows = list(reversed(topic_rows))
    labels = [r["topic"].replace("_", " ").title() for r in rows]
    values = [r["handoff_chunks"] for r in rows]
    width, height = 1500, 1180
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    font_path = Path(r"C:\Windows\Fonts\arial.ttf")
    bold_path = Path(r"C:\Windows\Fonts\arialbd.ttf")
    label_font = ImageFont.truetype(str(font_path), 27)
    value_font = ImageFont.truetype(str(bold_path), 26)
    axis_font = ImageFont.truetype(str(font_path), 24)
    left, right, top, bottom = 410, 100, 45, 80
    chart_w = width - left - right
    chart_h = height - top - bottom
    row_h = chart_h / len(rows)
    max_value = max(values)
    for tick in (0, 100, 200, 300):
        x = left + int(chart_w * tick / max_value)
        draw.line((x, top, x, height - bottom), fill="#E4E7EC", width=2)
        text = str(tick)
        box = draw.textbbox((0, 0), text, font=axis_font)
        draw.text((x - (box[2] - box[0]) / 2, height - bottom + 18), text, font=axis_font, fill="#667085")
    for idx, (label, value) in enumerate(zip(labels, values)):
        y_center = top + (idx + 0.5) * row_h
        box = draw.textbbox((0, 0), label, font=label_font)
        draw.text((left - 20 - (box[2] - box[0]), y_center - (box[3] - box[1]) / 2), label, font=label_font, fill="#344054")
        bar_h = int(row_h * 0.62)
        bar_w = int(chart_w * value / max_value)
        draw.rounded_rectangle((left, y_center - bar_h / 2, left + bar_w, y_center + bar_h / 2), radius=5, fill="#2E74B5")
        draw.text((left + bar_w + 12, y_center - 15), f"{value:,}", font=value_font, fill="#344054")
    title = "Chunks in handoff"
    box = draw.textbbox((0, 0), title, font=axis_font)
    draw.text((left + (chart_w - (box[2] - box[0])) / 2, height - 34), title, font=axis_font, fill="#475467")
    image.save(CHART_PATH, dpi=(180, 180))


def add_metric_strip(doc: Document, metrics: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(metrics))
    set_table_geometry(table, [2340] * len(metrics), indent_dxa=120)
    for cell, (number, label) in zip(table.rows[0].cells, metrics):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(number)
        set_run_font(r, size=17, bold=True, color=DARK_BLUE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(label)
        set_run_font(r2, size=8.5, bold=True, color=MID_GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_status_table(doc: Document) -> None:
    rows = [
        ("Dataset ID", "esg_3b1c0196c0f9"),
        ("Sectioning", "9,819 sections • contiguous_v2"),
        ("Chunking", "17,329 chunks • esg_chunk_v3"),
        ("Embedding context", "17,329 rows • esg_embed_ctx_v2"),
        ("Token bounds", "51–500 BGE tokens in the live dataset"),
        ("Citations", "17,329 exact page citations"),
        ("Vector manifest", "Rebuilt for all current chunks"),
        ("Vector database", "Not rebuilt yet"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2600, 6760])
    set_repeat_table_header(table.rows[0])
    for idx, title in enumerate(("Item", "Current state")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_paragraph_text(cell.paragraphs[0], title, size=10, bold=True, color=NAVY)
    for label, value in rows:
        cells = table.add_row().cells
        set_paragraph_text(cells[0].paragraphs[0], label, size=10, bold=True, color=DARK_BLUE)
        set_paragraph_text(cells[1].paragraphs[0], value, size=10, color=DARK)
    set_table_geometry(table, [2600, 6760])


def add_two_column_stats(doc: Document, rows: list[tuple[str, str]], headers=("Measure", "Value")) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [3900, 5460]
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, title in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_paragraph_text(cell.paragraphs[0], title, size=10, bold=True, color=NAVY)
    for label, value in rows:
        cells = table.add_row().cells
        set_paragraph_text(cells[0].paragraphs[0], label, size=9.5, bold=True, color=DARK_BLUE)
        set_paragraph_text(cells[1].paragraphs[0], value, size=9.5, color=DARK)
    set_table_geometry(table, widths)


def add_token_distribution_table(doc: Document, token_stats: dict, buckets: dict) -> None:
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    widths = [1200, 1300, 1500, 1360, 1300, 1400, 1300]
    set_table_geometry(table, widths)
    labels = ("Min", "P25", "Median", "Mean", "P75", "P90", "Max")
    values = (
        token_stats["min"], token_stats["p25"], token_stats["median"], token_stats["mean"],
        token_stats["p75"], token_stats["p90"], token_stats["max"],
    )
    for idx, label in enumerate(labels):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_paragraph_text(cell.paragraphs[0], label, size=9, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER)
    row = table.add_row()
    for idx, value in enumerate(values):
        set_paragraph_text(row.cells[idx].paragraphs[0], f"{value:,}" if isinstance(value, int) else str(value), size=9.5, bold=(idx == 2), color=DARK_BLUE if idx == 2 else DARK, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_table_geometry(table, widths)

    p = doc.add_paragraph("Token-size buckets")
    p.style = doc.styles["Heading 3"]
    bucket_rows = [(label + " tokens", f"{buckets.get(label, 0):,} chunks") for label in ("51–100", "101–250", "251–400", "401–500")]
    add_two_column_stats(doc, bucket_rows, headers=("Final BGE token count", "Chunks"))


def add_validation_table(doc: Document, checks: list[dict]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    widths = [4200, 1800, 1800, 1560]
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, title in enumerate(("Check", "Passed", "Expected", "Status")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_paragraph_text(cell.paragraphs[0], title, size=10, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT)
    for check in checks:
        cells = table.add_row().cells
        values = [check["check"], f'{check["passed_rows"]:,}', f'{check["expected_rows"]:,}', check["status"]]
        for idx, value in enumerate(values):
            set_paragraph_text(cells[idx].paragraphs[0], value, size=10, bold=(idx == 3), color=(GREEN if idx == 3 else DARK), align=WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)


def add_topic_table(doc: Document, rows: list[dict]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    widths = [5100, 2130, 2130]
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, title in enumerate(("ESG topic", "Handoff", "Safe pool")):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        set_paragraph_text(cell.paragraphs[0], title, size=9.5, bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT)
    for row in rows:
        cells = table.add_row().cells
        topic = row["topic"].replace("_", " ").title()
        values = (topic, f'{row["handoff_chunks"]:,}', f'{row["safe_pool_chunks"]:,}')
        for idx, value in enumerate(values):
            set_paragraph_text(cells[idx].paragraphs[0], value, size=9, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER if idx else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)


def add_note_box(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + " ")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=DARK)


def build() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    data = json.loads(DATA_PATH.read_text(encoding="utf-8"))
    validation = json.loads(VALIDATION_PATH.read_text(encoding="utf-8"))
    stats = json.loads(STATS_PATH.read_text(encoding="utf-8"))
    live_stats = stats["live_dataset"]
    handoff_stats = stats["handoff"]
    topic_rows = data["topic_counts"]
    checks = data["validation_checks"]
    make_chart(topic_rows)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)
    add_header_footer(section)

    # First-page masthead.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("TECHNICAL REPORT")
    set_run_font(r, size=10, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("ESG Chunk Dataset and\n2,000-Chunk Handoff")
    set_run_font(r, size=27, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("Current build status, sample design, validation, and review guide")
    set_run_font(r, size=13, color=MID_GRAY)

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [1800, 7560])
    for i, (label, value) in enumerate((
        ("Dataset", "esg_3b1c0196c0f9"),
        ("Chunk version", "esg_chunk_v3"),
        ("Report date", "31 July 2026"),
        ("Audience", "Technical review team"),
    )):
        set_paragraph_text(meta.cell(i, 0).paragraphs[0], label, size=10, bold=True, color=DARK_BLUE)
        set_paragraph_text(meta.cell(i, 1).paragraphs[0], value, size=10, color=DARK)
    set_table_geometry(meta, [1800, 7560])

    doc.add_paragraph()
    add_metric_strip(doc, [("17,329", "CURRENT CHUNKS"), ("2,000", "HANDOFF ROWS"), ("18", "ESG TOPICS"), ("0", "VALIDATION ERRORS")])

    doc.add_heading("Technical summary", level=1)
    doc.add_paragraph(
        "The new chunk dataset is built and internally aligned. It uses contiguous_v2 sectioning, "
        "esg_chunk_v3 chunking, and esg_embed_ctx_v2 embedding context. Every live chunk has an exact "
        "page citation. The vector manifest matches the current chunk set. The vector database has not "
        "been rebuilt, so search will not use this build until that final load is run."
    )
    add_note_box(
        doc,
        "Bottom line:",
        "The text, sections, chunks, citations, embedding inputs, and manifest are current. The remaining system step is rebuilding the vector database after human review or approval.",
    )

    doc.add_heading("Current dataset status", level=1)
    add_status_table(doc)

    doc.add_page_break()
    doc.add_heading("Live chunk statistics", level=1)
    doc.add_paragraph(
        "These figures describe all 17,329 current chunks. Token counts use the final BGE embedding input, including the metadata prefix and table context when present."
    )
    add_metric_strip(
        doc,
        [
            (f'{live_stats["tokens"]["median"]:,}', "MEDIAN TOKENS"),
            (f'{live_stats["page_spans"]["1 page"]:,}', "ONE-PAGE CHUNKS"),
            (f'{live_stats["multi_page"]:,}', "MULTI-PAGE CHUNKS"),
            (f'{live_stats["quality_tiers"]["narrative"]:,}', "NARRATIVE CHUNKS"),
        ],
    )

    doc.add_heading("Token length", level=2)
    add_token_distribution_table(doc, live_stats["tokens"], live_stats["token_buckets"])

    doc.add_heading("Structure and page span", level=2)
    overlap = live_stats["source_span_overlap"]
    add_two_column_stats(
        doc,
        [
            ("Reports / companies", f'{live_stats["documents"]:,} reports / {live_stats["companies"]:,} companies'),
            ("Sections", f'{live_stats["sections"]:,}'),
            ("Chunks per report", f'Median {live_stats["chunks_per_document"]["median"]:,}; P90 {live_stats["chunks_per_document"]["p90"]:,}; max {live_stats["chunks_per_document"]["max"]:,}'),
            ("Sections with one chunk", f'{live_stats["chunks_per_section"]["one_chunk_sections"]:,}'),
            ("Sections with more than one chunk", f'{live_stats["chunks_per_section"]["multi_chunk_sections"]:,}'),
            ("Page spans", f'{live_stats["page_spans"]["1 page"]:,} one-page; {live_stats["page_spans"]["2 pages"]:,} two-page; {live_stats["page_spans"]["3+ pages"]:,} three-or-more-page'),
            ("Adjacent chunk pairs with source overlap", f'{overlap["overlapping_pairs"]:,} of {overlap["adjacent_pairs"]:,} pairs'),
            ("Median chunk characters", f'{live_stats["chars"]["median"]:,}; P90 {live_stats["chars"]["p90"]:,}; max {live_stats["chars"]["max"]:,}'),
        ],
    )

    doc.add_heading("Chunk type and review tier", level=2)
    add_two_column_stats(
        doc,
        [
            ("Normal chunks", f'{live_stats["chunk_types"]["normal"]:,}'),
            ("Short evidence chunks", f'{live_stats["chunk_types"]["short_evidence"]:,}'),
            ("Narrative", f'{live_stats["quality_tiers"]["narrative"]:,}'),
            ("Layout-sensitive", f'{live_stats["quality_tiers"]["layout_sensitive"]:,}'),
            ("Noise", f'{live_stats["quality_tiers"]["noise"]:,}'),
            ("Exact citation status", f'{live_stats["citation_status"]["verified_exact"]:,} verified exact'),
        ],
    )

    doc.add_heading("Handoff scope", level=1)
    doc.add_paragraph(
        "The handoff contains 2,000 unique chunks from 188 reports and 105 companies. It covers all 18 ESG topics. "
        "The sample includes 1,715 narrative chunks and 285 layout-sensitive chunks. Reporting years are 2023 "
        "(1,041 chunks), 2024 (944), and 2025 (15)."
    )
    scope = [
        ("Token range", f'{validation["token_min"]}–{validation["token_max"]} BGE tokens; median {validation["token_median"]}'),
        ("Token-size buckets", f'{handoff_stats["token_buckets"]["51–100"]:,} at 51–100; {handoff_stats["token_buckets"]["101–250"]:,} at 101–250; {handoff_stats["token_buckets"]["251–400"]:,} at 251–400; {handoff_stats["token_buckets"]["401–500"]:,} at 401–500'),
        ("Page span", f'{validation["multi_page_chunks"]:,} chunks span more than one page'),
        ("Report cap", f'No report contributes more than {validation["max_chunks_per_document"]} chunks'),
        ("Exact citations", f'{validation["exact_citations"]:,} of {validation["rows"]:,}'),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2500, 6860])
    set_cell_shading(table.cell(0, 0), LIGHT_GRAY)
    set_cell_shading(table.cell(0, 1), LIGHT_GRAY)
    set_paragraph_text(table.cell(0, 0).paragraphs[0], "Measure", size=10, bold=True, color=NAVY)
    set_paragraph_text(table.cell(0, 1).paragraphs[0], "Value", size=10, bold=True, color=NAVY)
    for label, value in scope:
        cells = table.add_row().cells
        set_paragraph_text(cells[0].paragraphs[0], label, size=10, bold=True, color=DARK_BLUE)
        set_paragraph_text(cells[1].paragraphs[0], value, size=10, color=DARK)
    set_table_geometry(table, [2500, 6860])

    doc.add_page_break()
    doc.add_heading("Topic distribution", level=1)
    doc.add_paragraph(
        "The sample gives every ESG topic at least 30 chunks, then allocates the remaining rows in line with topic size. "
        "Human capital has the most rows, followed by supply-chain ethics and environmental content."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(CHART_PATH), width=Inches(6.25))
    chart_shape = doc.inline_shapes[-1]
    chart_shape._inline.docPr.set(
        "descr",
        "Horizontal bar chart showing the count of handoff chunks for each of 18 ESG topics; human capital is largest at 304 chunks.",
    )
    caption = doc.add_paragraph("Figure 1. Chunks in the 2,000-row handoff by ESG topic.")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    for run in caption.runs:
        set_run_font(run, size=9, italic=True, color=MID_GRAY)

    doc.add_page_break()
    doc.add_heading("Topic allocation detail", level=1)
    doc.add_paragraph("Exact counts for the handoff and the safe pool used during sampling.")
    add_topic_table(doc, topic_rows)

    doc.add_page_break()
    doc.add_heading("Validation results", level=1)
    doc.add_paragraph(
        "All five row-level checks passed for every selected chunk. The checks compare the package with current chunk files, "
        "parsed source spans, embedding context, and manifest state. There are no duplicate chunk IDs and no reported validation errors."
    )
    add_validation_table(doc, checks)
    doc.add_paragraph()
    add_note_box(
        doc,
        "What this proves:",
        "The package is internally consistent and traceable to parsed source text. Human review is still needed to judge reading order, sentence boundaries, metadata meaning, and table usability.",
    )

    doc.add_heading("Sampling method", level=1)
    for text in (
        "Use deterministic seed 20260731.",
        "Give each topic a minimum of 30 chunks.",
        "Allocate remaining rows in line with topic size.",
        "Keep the narrative and layout-sensitive mix inside each topic.",
        "Limit each report to 15 chunks.",
        "Exclude noise and held rows.",
    ):
        doc.add_paragraph(text, style="List Bullet")

    doc.add_heading("Limits", level=1)
    doc.add_paragraph(
        "This is a designed review sample, not a simple random sample. Review results should not be used as a raw statistical "
        "defect rate. Some section titles are table fragments. Reviewers should focus on whether chunk_text is readable and can "
        "support a cited answer. Layout-sensitive chunks may place a number farther from its label than the source page did; flag it "
        "only when meaning is lost or changed."
    )

    doc.add_heading("Review instructions", level=1)
    for text in (
        "Open review_template.csv.",
        "Read chunk_text and its metadata.",
        "Set review_status to pass or fail.",
        "For a failure, add a short issue_type such as order, cutoff, metadata, or table.",
        "Add a short note that explains the problem.",
        "Return both chunk_id and chunk_text_sha256.",
    ):
        doc.add_paragraph(text, style="List Number")
    add_note_box(
        doc,
        "Important:",
        "chunk_id can change after a rebuild. chunk_text_sha256 is the stable text fingerprint. Returning both fields lets earlier review work carry forward when the text is unchanged.",
    )

    doc.core_properties.title = "ESG Chunk Dataset and 2,000-Chunk Handoff"
    doc.core_properties.subject = "Technical status and review handoff"
    doc.core_properties.author = "Retail Intelligence ESG"
    doc.core_properties.keywords = "ESG, chunks, embeddings, validation, handoff"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
